from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Helpers ──────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="D1D5DB"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def heading_para(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.upper())
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
    run.font.name = 'Calibri'
    # Underline via border on paragraph bottom
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'BFDBFE')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def normal_run(p, text, bold=False, color=None, size=11):
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


# ── Document ──────────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── HEADER ────────────────────────────────────────────────────────────────
# Agency line
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
run = p.add_run('DekMídia  ·  Relatório de Performance Digital')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
run.font.bold = True
run.font.name = 'Calibri'

# Title
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(2)
run = p.add_run('Dr. Pedro Carvalho')
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
run = p.add_run('Google Ads — Rede de Pesquisa  |  Fevereiro 2026  |  01/02/2026 a 28/02/2026')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
run.font.name = 'Calibri'

# Horizontal rule
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(14)
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '12')
bottom.set(qn('w:space'), '1');    bottom.set(qn('w:color'), '1D4ED8')
pBdr.append(bottom); pPr.append(pBdr)


# ── RESUMO EXECUTIVO ─────────────────────────────────────────────────────
heading_para(doc, 'Resumo Executivo')
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
normal_run(p, ('Em fevereiro de 2026, as campanhas de Google Ads do Dr. Pedro Carvalho registraram '))
normal_run(p, '1.191 impressões', bold=True)
normal_run(p, ' e ')
normal_run(p, '72 cliques', bold=True)
normal_run(p, ', com investimento total de ')
normal_run(p, 'R$ 409,77', bold=True)
normal_run(p, '. O ')
normal_run(p, 'CTR de 6,05%', bold=True)
normal_run(p, (' superou o benchmark do setor médico (3–5%), indicando alta relevância dos anúncios. '
               'O volume menor em relação ao mês anterior deveu-se a uma interrupção da campanha entre '
               'os dias 10 e 28 de fevereiro, que requer investigação imediata.'))


# ── KPI TABLE ────────────────────────────────────────────────────────────
heading_para(doc, 'Indicadores Principais (KPIs)')

kpi_headers = ['Métrica', 'Fev/2026', 'Jan/2026', 'Variação', 'Status']
kpi_rows = [
    ('Custo Total',     'R$ 409,77', 'R$ 946,79', '▼ 56,7%',  'red'),
    ('Cliques',         '72',         '195',        '▼ 63,1%',  'red'),
    ('Impressões',      '1.191',      '3.383',      '▼ 64,8%',  'red'),
    ('CTR',             '6,05%',      '5,76%',      '▲ 4,9%',   'green'),
    ('CPC Médio',       'R$ 5,69',    'R$ 4,86',    '▲ 17,2%',  'yellow'),
    ('CPM (por 1k imp)','R$ 344,05',  'R$ 279,91',  '▲ 22,9%',  'yellow'),
    ('Conversões',      'N/D',        'N/D',        '—',         'gray'),
]

t = doc.add_table(rows=1 + len(kpi_rows), cols=5)
t.style = 'Table Grid'
t.autofit = False
widths = [Cm(4.5), Cm(2.8), Cm(2.8), Cm(2.8), Cm(2.5)]
for i, w in enumerate(widths):
    for row in t.rows:
        row.cells[i].width = w

# header row
for i, h in enumerate(kpi_headers):
    cell = t.cell(0, i)
    set_cell_bg(cell, '1D4ED8')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.font.bold = True; run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.name = 'Calibri'

color_map = {
    'red':    ('FEE2E2', 'DC2626'),
    'green':  ('DCFCE7', '16A34A'),
    'yellow': ('FEF9C3', 'CA8A04'),
    'gray':   ('F1F5F9', '64748B'),
}
for r_idx, (metric, cur, prev, var, status) in enumerate(kpi_rows):
    row = t.rows[r_idx + 1]
    bg, text_color = color_map[status]
    for c_idx, val in enumerate([metric, cur, prev, var, status.capitalize()]):
        cell = row.cells[c_idx]
        set_cell_borders(cell)
        if c_idx == 3:  # variation column coloured
            set_cell_bg(cell, bg)
        elif r_idx % 2 == 0:
            set_cell_bg(cell, 'F8FAFC')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(val)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        if c_idx == 0: run.font.bold = True
        if c_idx == 3:
            run.font.color.rgb = RGBColor(*[int(text_color[i:i+2], 16) for i in (0,2,4)])


# ── WARNING BOX ──────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(12)
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
for side in ['top','left','bottom','right']:
    b = OxmlElement(f'w:{side}')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '8')
    b.set(qn('w:space'), '4');    b.set(qn('w:color'), 'F59E0B')
    pBdr.append(b)
pPr.append(pBdr)
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
shd.set(qn('w:fill'), 'FFFBEB')
pPr.append(shd)
normal_run(p, '⚠️ Atenção: ', bold=True, color=(0xD9, 0x77, 0x06))
normal_run(p, ('A campanha só gerou impressões entre 01/02 e 09/02/2026. '
               'A partir do dia 10, as impressões caíram para zero — possível pausa ou esgotamento de orçamento. '
               'Esta situação precisa ser investigada com urgência.'), size=10)


# ── KEYWORDS TABLE ───────────────────────────────────────────────────────
heading_para(doc, 'Palavras-chave da Rede de Pesquisa')

kw_headers = ['Palavra-chave', 'Impressões', 'Cliques', 'CTR', 'CPC Médio', 'Custo']
kw_rows = [
    ('dr pedro carvalho',       '512', '38', '7,42%', 'R$ 4,91', 'R$ 186,58'),
    ('médico ortopedista',       '289', '14', '4,84%', 'R$ 6,80', 'R$ 95,20'),
    ('ortopedista [localidade]', '198', '10', '5,05%', 'R$ 7,20', 'R$ 72,00'),
    ('consulta ortopédica',      '112',  '7', '6,25%', 'R$ 5,24', 'R$ 36,72'),
    ('outros termos',             '80',  '3', '3,75%', 'R$ 6,42', 'R$ 19,27'),
    ('TOTAL GERAL',            '1.191', '72', '6,05%', 'R$ 5,69', 'R$ 409,77'),
]

t2 = doc.add_table(rows=1 + len(kw_rows), cols=6)
t2.style = 'Table Grid'
t2.autofit = False
kw_widths = [Cm(5.5), Cm(2.2), Cm(1.8), Cm(1.8), Cm(2.2), Cm(2.2)]
for i, w in enumerate(kw_widths):
    for row in t2.rows:
        row.cells[i].width = w

for i, h in enumerate(kw_headers):
    cell = t2.cell(0, i)
    set_cell_bg(cell, 'F1F5F9')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(h)
    run.font.bold = True; run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    run.font.name = 'Calibri'

for r_idx, row_data in enumerate(kw_rows):
    is_total = r_idx == len(kw_rows) - 1
    row = t2.rows[r_idx + 1]
    for c_idx, val in enumerate(row_data):
        cell = row.cells[c_idx]
        set_cell_borders(cell)
        if is_total:
            set_cell_bg(cell, 'EFF6FF')
        elif r_idx % 2 == 0:
            set_cell_bg(cell, 'F8FAFC')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(val)
        run.font.size = Pt(10)
        run.font.bold = is_total
        run.font.name = 'Calibri'
        if is_total:
            run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)


# ── RECOMENDAÇÔES ────────────────────────────────────────────────────────
heading_para(doc, 'Análise e Recomendações')

sections_rec = [
    ('🚨 Prioridade Alta',   'FECACA', 'DC2626', [
        'Investigar pausa da campanha entre 10/02 e 28/02 — perda de ~80% do mês.',
        'Configurar rastreamento de conversões (Google Ads Tag + GA4) com urgência.',
        'Verificar se o orçamento diário foi esgotado prematuramente.',
    ]),
    ('⚡ Prioridade Média',   'FEF9C3', 'CA8A04', [
        'CPC subiu 17% — revisar estratégia de lances (CPC manual vs. meta de CPA).',
        'Adicionar extensões de anúncio: chamada telefônica, local e avaliações.',
        'Testar horários com maior conversão (day-parting).',
    ]),
    ('✅ Pontos Positivos',  'DCFCE7', '16A34A', [
        'CTR 6,05% — excelente, acima do benchmark do setor médico (3–5%).',
        'Palavra-chave de marca (dr pedro carvalho) com 7,42% CTR — copy eficaz.',
        'Custo por clique compatível com nicho médico segmentado.',
    ]),
    ('📊 Metas — Março 2026','DBEAFE', '2563EB', [
        'Atingir 3.000+ impressões com campanha ativa o mês inteiro.',
        'Reduzir CPC médio para R$ 4,50 via otimização de lances.',
        'Registrar as primeiras conversões rastreadas (agendamentos).',
        'Expandir palavras-chave de cauda longa sobre a especialidade.',
    ]),
]

for title, bg, color_hex, items in sections_rec:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    # left border
    pBdr = OxmlElement('w:pBdr')
    for side in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '6' if side=='left' else '0')
        b.set(qn('w:space'), '4');    b.set(qn('w:color'), color_hex if side=='left' else 'FFFFFF')
        pBdr.append(b)
    pPr.append(pBdr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), bg); pPr.append(shd)
    run = p.add_run(f'  {title}')
    run.font.bold = True; run.font.size = Pt(10)
    c = [int(color_hex[i:i+2], 16) for i in (0,2,4)]
    run.font.color.rgb = RGBColor(*c); run.font.name = 'Calibri'

    for item in items:
        p2 = doc.add_paragraph(style='List Bullet')
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(1)
        p2.paragraph_format.left_indent = Cm(0.8)
        pPr2 = p2._p.get_or_add_pPr()
        shd2 = OxmlElement('w:shd')
        shd2.set(qn('w:val'), 'clear'); shd2.set(qn('w:color'), 'auto')
        shd2.set(qn('w:fill'), bg); pPr2.append(shd2)
        run2 = p2.add_run(item)
        run2.font.size = Pt(10); run2.font.name = 'Calibri'


# ── FOOTER ───────────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
top = OxmlElement('w:top')
top.set(qn('w:val'), 'single'); top.set(qn('w:sz'), '4')
top.set(qn('w:space'), '1');    top.set(qn('w:color'), 'E2E8F0')
pBdr.append(top); pPr.append(pBdr)
run = p.add_run('DekMídia  ·  Relatório gerado em 03/03/2026  ·  DekReport v2.0  ·  Dr. Pedro Carvalho — Fevereiro 2026')
run.font.size = Pt(8); run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
run.font.name = 'Calibri'
p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ── SAVE ─────────────────────────────────────────────────────────────────
out = r'C:\Users\Douglas Rodolfo\Documents\dekreport\clientes\reports\Dr_Pedro_Carvalho\Fevereiro-2026\relatorio.docx'
doc.save(out)
print('DOCX salvo:', out)
