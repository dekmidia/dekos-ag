from docx import Document
import os

def create_msg_doc(id_lead, empresa, canal, msg):
    doc = Document()
    doc.add_heading(f'Abordagem: {empresa}', 0)
    doc.add_paragraph(f'ID Lead: {id_lead}')
    doc.add_paragraph(f'Canal Recomendado: {canal}')
    doc.add_paragraph('---')
    doc.add_paragraph(msg)
    
    output_dir = 'c:/Users/crist/Documents/Dev/ANTIGRAVITY/dekos-ag/PROJETOS/_prospection/mensagens-prontas'
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    filename = f'{id_lead}_msg.docx'
    doc.save(os.path.join(output_dir, filename))
    print(f'Documento salvo: {filename}')

# Mensagens aprovadas
mensagens = [
    {
        "id": "20260310-AESTHETIC-CPQ-001",
        "empresa": "Grupo Vip Rejuveness",
        "canal": "WhatsApp",
        "msg": "Reparei que a Rejuveness tem uma avaliação impecável (4.9 estrelas) em Campinas - parabéns pelo trabalho.\nNotei que, apesar da autoridade física, vocês não aparecem nos anúncios patrocinados quando buscamos por 'estética avançada' na região hoje.\nIsso costuma deixar muito espaço livre para concorrentes menores captarem leads que deveriam ser seus.\nFaz sentido conversarmos 15 min sobre como escalar sua agenda? [Seu Nome | DekMidia]"
    },
    {
        "id": "20260310-AESTHETIC-IND-001",
        "empresa": "Clínica Vitessi",
        "canal": "WhatsApp",
        "msg": "Notei a estrutura tecnológica de ponta da Vitessi em Indaiatuba - impressionante o complexo de 600m².\nReparei que, apesar de terem um site robusto, o alcance orgânico quando buscamos por 'estética avançada indaiatuba' ainda está muito abaixo do potencial do negócio.\nIsso faz com que o custo por lead de vocês seja maior do que poderia, já que dependem 100% de indicações ou tráfego pago instável.\nJá ajudamos clínicas a triplicar o tráfego orgânico com SEO focado em procedimentos de alta margem. Faz sentido conversarmos 15 min? [Seu Nome | DekMidia]"
    },
    {
        "id": "20260310-REALESTATE-BC-001",
        "empresa": "Guilherme Pilger Corretor",
        "canal": "WhatsApp",
        "msg": "Guilherme, acompanho seu posicionamento em Balneário e a autoridade que você construiu no mercado de luxo é referência.\nVi que o fluxo de leads de alto padrão em BC está no topo, mas muitos corretores perdem escala por falta de uma triagem automatizada eficiente antes do contato pessoal.\nUma estrutura de landing page específica para lançamentos de luxo poderia filtrar 3x mais clientes qualificados para você.\nPodemos falar 10 min sobre automatizar essa triagem? [Seu Nome | DekMidia]"
    },
    {
        "id": "20260310-REALESTATE-ALPH-001",
        "empresa": "Gisele Alphaville Imóveis",
        "canal": "WhatsApp",
        "msg": "Gisele, vi sua presença em Alphaville e o padrão dos imóveis que você comercializa é altíssimo.\nNotei um ponto técnico: seu site atual está levando mais de 5 segundos para carregar no mobile. No mercado de luxo, 1 segundo a mais de espera faz o cliente desistir e ir para o concorrente no Instagram.\nUma Landing Page 'Ultra-Fast' focada em conversão para seus residenciais de destaque poderia acelerar suas vendas este mês.\nPodemos falar 10 min sobre como otimizar essa velocidade de conversão? [Seu Nome | DekMidia]"
    },
    {
        "id": "20260310-LAW-HIGH-001",
        "empresa": "Cescon Barrieu",
        "canal": "LinkedIn",
        "msg": "Olá [Sócio], acompanho a trajetória da Cescon Barrieu em M&A e o impacto de vocês no mercado brasileiro é notável.\nNotei que, para uma firma desse porte, a apresentação digital de propostas e o branding nas redes pode ser ainda mais dinâmico para atrair clientes globais que buscam agilidade.\nImplementamos um design de propostas premium que tem ajudado firmas similares a fechar contratos 20% mais rápido por puro impacto visual e clareza.\nFaz sentido me conectar para eu te mostrar esse case? [Seu Nome | DekMidia]"
    }
]

for m in mensagens:
    create_msg_doc(m['id'], m['empresa'], m['canal'], m['msg'])
